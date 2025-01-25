import streamlit as st
import pandas as pd
import numpy as np
from scipy.stats import ttest_ind, f_oneway
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
from docx import Document


# Function to create a Word document
def create_word_doc(content):
    doc = Document()
    for section in content:
        doc.add_heading(section['title'], level=1)
        for table in section.get('tables', []):
            doc.add_paragraph(f"Table: {table['title']}")
            df = table['dataframe']
            table_doc = doc.add_table(rows=1, cols=len(df.columns))
            table_doc.style = 'Table Grid'
            # Add headers
            hdr_cells = table_doc.rows[0].cells
            for i, col in enumerate(df.columns):
                hdr_cells[i].text = str(col)
            # Add rows
            for _, row in df.iterrows():
                row_cells = table_doc.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
        for chart in section.get('charts', []):
            doc.add_paragraph(f"Chart: {chart['title']}")
            image_stream = chart["image_buffer"]
            doc.add_picture(image_stream)
    return doc


# Helper function to apply manual ranges
def apply_manual_ranges(value, ranges):
    sorted_ranges = sorted(ranges, key=lambda x: float(x[1:]) if x[0] in "<>" else float(x.split("-")[0]))
    for r in sorted_ranges:
        if r.startswith("<"):
            threshold = float(r[1:])
            if value < threshold:
                return r
        elif r.startswith(">"):
            threshold = float(r[1:])
            if value > threshold:
                return r
        elif "-" in r:
            low, high = map(float, r.split("-"))
            if low <= value < high:
                return r
    return "Other"  # Catch-all for values outside specified ranges


# Upload data
st.title("Eternals Thesis")
uploaded_file = st.file_uploader("Upload your Excel file", type=["xlsx"])

if uploaded_file:
    df = pd.read_excel(uploaded_file)
    export_content = []

    # Tab structure
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["Distribution Tables", "Pivot Tables", "Statistical Analysis", "Correlations", "Graph Builder"])

    # Tab 1: Distribution Tables
    with tab1:
        st.header("Automated Distribution Tables")
        tab1_content = {"title": "Distribution Tables", "tables": [], "charts": []}

        # Combined Distribution Table
        combined_columns = st.multiselect("Select Columns for Combined Distribution", df.columns)
        if combined_columns:
            st.write(f"Combined Distribution for Columns: {', '.join(combined_columns)}")
            combined_distribution = df[combined_columns].apply(lambda row: tuple(row), axis=1).value_counts().reset_index()
            combined_distribution.columns = ["Combination", "Count"]
            combined_distribution["Percentage"] = (combined_distribution["Count"] / combined_distribution["Count"].sum() * 100).round(2).astype(str) + "%"
            combined_distribution.reset_index(drop=True, inplace=True)
            combined_distribution.index = combined_distribution.index + 1  # Start index from 1

            # Convert all columns to strings
            combined_distribution = combined_distribution.astype(str)

            st.dataframe(combined_distribution)
            tab1_content["tables"].append({"title": "Combined Distribution", "dataframe": combined_distribution})

        # Individual Column Distribution
        for column in df.columns:
            if df[column].dtype in [np.int64, np.float64, object]:
                st.subheader(f"Distribution for {column}")

                # Option to use manual ranges or automatic binning
                use_manual_ranges = st.checkbox(f"Use Manual Ranges for {column}?", key=f"{column}_manual_ranges")
                manual_ranges = []
                if use_manual_ranges and df[column].dtype in [np.int64, np.float64]:
                    st.write("Specify manual ranges (e.g., '<5', '5-10', '>90')")
                    manual_ranges = st.text_area(
                        f"Enter ranges for {column} (one range per line)", 
                        value="<5\n5-10\n10-20\n>90",
                        key=f"{column}_manual_range_input"
                    ).splitlines()

                # Use automatic binning if manual ranges are not specified
                if not use_manual_ranges or not manual_ranges:
                    use_ranges = st.checkbox(f"Use Dynamic Ranges for {column}?", key=f"{column}_ranges")
                    if use_ranges and df[column].dtype in [np.int64, np.float64]:
                        range_step = st.number_input(
                            f"Step size for {column} ranges",
                            min_value=0.01 if df[column].dtype == np.float64 else 1,
                            value=10 if df[column].dtype == np.int64 else 0.1,
                            key=f"{column}_range_step",
                        )
                        bins = np.arange(df[column].min(), df[column].max() + range_step, range_step)
                        labels = [f"{round(bins[i], 2)}-{round(bins[i + 1], 2)}" for i in range(len(bins) - 1)]
                        df[column] = pd.cut(df[column], bins=bins, labels=labels, right=False)
                elif use_manual_ranges:
                    df[column] = df[column].apply(lambda x: apply_manual_ranges(x, manual_ranges))

                # Create distribution table
                distribution = df[column].value_counts().reset_index()
                distribution.columns = [column, "Count"]
                distribution["Percentage"] = (distribution["Count"] / distribution["Count"].sum() * 100).round(2).astype(str) + '%'
                distribution.reset_index(drop=True, inplace=True)
                distribution.index = distribution.index + 1  # Start index from 1

                if not distribution.empty:
                    total_row = pd.DataFrame({column: ["Total"], "Count": [distribution["Count"].sum()], "Percentage": ["100%"]})
                    distribution = pd.concat([distribution, total_row], ignore_index=True)

                    st.dataframe(distribution)
                    tab1_content["tables"].append({"title": f"Distribution for {column}", "dataframe": distribution})

                    # Graph Customization Options
                    graph_title = st.text_input(f"Graph Title for {column}", value=f"{column} Distribution", key=f"{column}_title")
                    x_label = st.text_input(f"X-Axis Label for {column}", value=column, key=f"{column}_x_label")
                    y_label = st.text_input(f"Y-Axis Label for {column}", value="Count", key=f"{column}_y_label")
                    legend_label = st.text_input(f"Legend Label for {column}", value="Values", key=f"{column}_legend")
                    x_axis_orientation = st.radio(
                        f"X-Axis Label Orientation for {column}",
                        options=["Horizontal", "Vertical"],
                        index=0,
                        key=f"{column}_orientation"
                    )

                    # Plot chart with color differentiation
                    fig, ax = plt.subplots()
                    colors = sns.color_palette("Set2", len(distribution.iloc[:-1]))
                    distribution.iloc[:-1].plot(kind="bar", x=column, y="Count", ax=ax, legend=False, color=colors)
                    ax.set_title(graph_title)
                    ax.set_xlabel(x_label)
                    ax.set_ylabel(y_label)
                    ax.legend([legend_label])
                    rotation_angle = 0 if x_axis_orientation == "Horizontal" else 90
                    ax.tick_params(axis="x", rotation=rotation_angle)
                    st.pyplot(fig)

                    # Save chart for Word export
                    buffer = BytesIO()
                    fig.savefig(buffer, format="png")
                    buffer.seek(0)
                    tab1_content["charts"].append({"title": f"{column} Distribution", "image_buffer": buffer})
                else:
                    st.info(f"No data available for column {column}.")

        export_content.append(tab1_content)

    # Tab 2: Pivot Tables
    with tab2:
        st.header("Pivot Tables")
        st.write("Select columns to create pivot tables.")
        rows = st.multiselect("Rows", df.columns)
        cols = st.multiselect("Columns", df.columns)
        values = st.selectbox("Values", df.columns)
        agg_func = st.selectbox("Aggregation Function", ["mean", "sum", "count", "max", "min"])
        filters = st.multiselect("Filters", df.columns)

        if st.button("Generate Pivot Table"):
            try:
                pivot_table = pd.pivot_table(df, index=rows, columns=cols, values=values, aggfunc=agg_func)
                pivot_table.reset_index(drop=True, inplace=True)
                pivot_table.index = pivot_table.index + 1  # Start index from 1
                st.dataframe(pivot_table)
            except Exception as e:
                st.error(f"Error generating pivot table: {e}")

    # Tab 3: Statistical Analysis
    with tab3:
        st.header("Statistical Analysis")
        st.write("Select columns for statistical calculations.")
        selected_columns = st.multiselect("Select Columns", df.select_dtypes(include=[np.number]).columns)
        
        # T-Test
        if len(selected_columns) == 2:
            col1, col2 = selected_columns[:2]
            st.write(f"Calculating statistics between {col1} and {col2}")

            stats = {
                "Metric": ["Mean", "Median", "Std Dev", "T-Statistic", "P-Value"],
                col1: [
                    df[col1].mean(),
                    df[col1].median(),
                    df[col1].std(),
                    None,  # Placeholder for T-Statistic
                    None,  # Placeholder for P-Value,
                ],
                col2: [
                    df[col2].mean(),
                    df[col2].median(),
                    df[col2].std(),
                    None,
                    None,
                ],
            }
            t_stat, p_value = ttest_ind(df[col1].dropna(), df[col2].dropna())
            stats[col1][3] = t_stat
            stats[col1][4] = p_value
            stats[col2][3] = t_stat
            stats[col2][4] = p_value

            stats_df = pd.DataFrame(stats)
            stats_df.index = stats_df.index + 1  # Start index from 1
            st.dataframe(stats_df)

        # ANOVA
        if len(selected_columns) > 2:
            st.write("Performing ANOVA test for selected columns.")
            groups = [df[col].dropna() for col in selected_columns]
            f_stat, p_value = f_oneway(*groups)
            st.write(f"ANOVA F-Statistic: {f_stat:.4f}")
            st.write(f"ANOVA P-Value: {p_value:.4f}")

    # Tab 4: Correlations
    with tab4:
        st.header("Correlations")
        st.write("Correlation matrix of numeric columns.")
        numeric_df = df.select_dtypes(include=[np.number])
        if numeric_df.empty:
            st.warning("No numeric columns available for correlation.")
        else:
            correlation_matrix = numeric_df.corr()
            correlation_matrix.reset_index(drop=True, inplace=True)
            correlation_matrix.index = correlation_matrix.index + 1  # Start index from 1
            st.dataframe(correlation_matrix)

            fig, ax = plt.subplots()
            sns.heatmap(correlation_matrix, annot=True, cmap="coolwarm", ax=ax)
            ax.set_title("Correlation Heatmap")
            st.pyplot(fig)

    # Tab 5: Graph Builder
    with tab5:
        st.header("Graph Builder")
        st.write("Select columns to build graphs.")
        x_col = st.selectbox("X-Axis", df.columns)
        y_col = st.selectbox("Y-Axis", df.columns)
        graph_type = st.selectbox("Graph Type", ["Scatter", "Line", "Bar", "Histogram", "Boxplot"])

        graph_title = st.text_input("Graph Title", value=f"{x_col} vs {y_col}")
        x_label = st.text_input("X-Axis Label", value=x_col)
        y_label = st.text_input("Y-Axis Label", value=y_col)

        if st.button("Generate Graph"):
            fig, ax = plt.subplots()
            if graph_type == "Scatter":
                sns.scatterplot(x=df[x_col], y=df[y_col], ax=ax)
            elif graph_type == "Line":
                sns.lineplot(x=df[x_col], y=df[y_col], ax=ax)
            elif graph_type == "Bar":
                sns.barplot(x=df[x_col], y=df[y_col], ax=ax)
            elif graph_type == "Histogram":
                sns.histplot(df[x_col], bins=30, kde=True, ax=ax)
            elif graph_type == "Boxplot":
                sns.boxplot(x=df[x_col], y=df[y_col], ax=ax)
            ax.set_title(graph_title)
            ax.set_xlabel(x_label)
            ax.set_ylabel(y_label)
            st.pyplot(fig)

    # Download Button
    if st.button("Download as Word Document"):
        doc = create_word_doc(export_content)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button("Download Word Document", buffer, "data_analysis.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
